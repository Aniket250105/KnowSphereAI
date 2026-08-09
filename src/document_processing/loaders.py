import pymupdf  # PyMuPDF
import docx
from pathlib import Path
from datetime import datetime, timezone
import os

from src.core.logger import get_logger
from src.models.document import Document, generate_document_id
from src.document_processing.base import BaseDocumentLoader

logger = get_logger(__name__)

def _get_file_stats(file_path: Path) -> dict:
    """Helper function to extract creation/modification time and ID."""
    try:
        stat = file_path.stat()
        created_at = datetime.fromtimestamp(stat.st_ctime, tz=timezone.utc).isoformat()
        modified_at = datetime.fromtimestamp(stat.st_mtime, tz=timezone.utc).isoformat()
    except Exception as e:
        logger.warning(f"Failed to extract file stats for {file_path}: {e}")
        created_at = ""
        modified_at = ""
        
    try:
        rel_path = str(file_path.relative_to(Path.cwd()))
    except ValueError:
        rel_path = str(file_path)

    return {
        "created_at": created_at,
        "modified_at": modified_at,
        "id": generate_document_id(file_path),
        "path": rel_path
    }

class PDFLoader(BaseDocumentLoader):
    """Loader for PDF files using PyMuPDF."""
    
    def load(self, file_path: Path) -> Document:
        logger.info(f"Loaded PDF {file_path.name}")
        text = ""
        page_count = 0
        try:
            with pymupdf.open(str(file_path)) as doc:
                page_count = len(doc)
                for page in doc:
                    text += page.get_text() + "\n"
        except Exception as e:
            logger.error(f"Error reading PDF file {file_path}: {e}")
            raise

        stats = _get_file_stats(file_path)
        
        return Document(
            name=file_path.name,
            file_type="pdf",
            content=text,
            path=stats["path"],
            id=stats["id"],
            pages=page_count,
            characters=len(text),
            created_at=stats["created_at"],
            modified_at=stats["modified_at"]
        )

class TXTLoader(BaseDocumentLoader):
    """Loader for plain text files."""
    
    def load(self, file_path: Path) -> Document:
        logger.info(f"Loaded TXT {file_path.name}")
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
        except UnicodeDecodeError:
            logger.warning(f"UTF-8 decode failed for {file_path}, falling back to latin-1")
            with open(file_path, "r", encoding="latin-1") as f:
                text = f.read()
        except Exception as e:
            logger.error(f"Error reading TXT file {file_path}: {e}")
            raise
            
        stats = _get_file_stats(file_path)
        
        return Document(
            name=file_path.name,
            file_type="txt",
            content=text,
            path=stats["path"],
            id=stats["id"],
            pages=1,
            characters=len(text),
            created_at=stats["created_at"],
            modified_at=stats["modified_at"]
        )

class DOCXLoader(BaseDocumentLoader):
    """Loader for Microsoft Word (DOCX) files using python-docx."""
    
    def load(self, file_path: Path) -> Document:
        logger.info(f"Loaded DOCX {file_path.name}")
        try:
            doc = docx.Document(str(file_path))
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        except Exception as e:
            logger.error(f"Error reading DOCX file {file_path}: {e}")
            raise
            
        stats = _get_file_stats(file_path)
        
        return Document(
            name=file_path.name,
            file_type="docx",
            content=text,
            path=stats["path"],
            id=stats["id"],
            pages=1,
            characters=len(text),
            created_at=stats["created_at"],
            modified_at=stats["modified_at"]
        )
